#!/usr/bin/env python3
"""Generate a Word catalog of paintings with QR codes and links."""

import base64
import io
import json
import os
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

SITE_DIR = os.path.join(os.path.dirname(__file__), 'site')
DATA_PATH = os.path.join(SITE_DIR, 'data.js')
QR_PATH = os.path.join(SITE_DIR, 'qr_data.js')
OUT_DOCX = os.path.join(SITE_DIR, 'gallery_catalog.docx')


def parse_js_object(text, key):
    pattern = re.compile(rf"window\.{re.escape(key)}\s*=\s*\{{(.*?)\}}\s*;", re.S)
    match = pattern.search(text)
    if not match:
        raise ValueError(f'Could not find window.{key} object in JS file')
    body = match.group(1)
    body = re.sub(r"([\w$]+)\s*:", r'"\1":', body)
    body = body.replace("'", '"')
    body = re.sub(r',\s*([\}\]])', r'\1', body)
    return body


def load_gallery_data():
    with open(DATA_PATH, encoding='utf-8') as f:
        text = f.read()

    body = parse_js_object(text, 'GALLERY_DATA')
    data = json.loads(f'{{{body}}}')
    paintings = data.get('paintings', [])
    base = data.get('base', '').rstrip('/')
    return base, paintings


def load_qr_data():
    with open(QR_PATH, encoding='utf-8') as f:
        text = f.read()

    body = parse_js_object(text, 'QR_DATA')
    qr = json.loads(f'{{{body}}}')
    return qr


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def build_document(base, paintings, qr_data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run('Приморская картинная галерея — каталог картин')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('Таблица содержит автора, название, ссылку на страницу и QR-код для каждой картины.', style='Intense Quote')
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Автор / Название'
    hdr_cells[2].text = 'Ссылка'
    hdr_cells[3].text = 'QR-код'

    for painting in paintings:
        row_cells = table.add_row().cells
        row_cells[0].text = str(painting.get('id', ''))

        author = painting.get('artist', '')
        life = painting.get('life', '')
        title = painting.get('title', '')
        author_line = author
        if life:
            author_line = f'{author} ({life})'

        row_cells[1].text = f'{author_line}\n{title}'
        paragraph = row_cells[2].paragraphs[0]
        link = base + '/#p=' + str(painting.get('id', ''))
        add_hyperlink(paragraph, link, link)

        img_data = qr_data.get(str(painting.get('id')))
        if img_data:
            image_bytes = base64.b64decode(img_data.split(',', 1)[1])
            image_stream = io.BytesIO(image_bytes)
            paragraph = row_cells[3].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(1.2))
        else:
            row_cells[3].text = 'QR отсутствует'

    doc.add_page_break()
    return doc


def main():
    base, paintings = load_gallery_data()
    qr_data = load_qr_data()
    doc = build_document(base, paintings, qr_data)
    doc.save(OUT_DOCX)
    print(f'Word catalog written to: {OUT_DOCX}')


if __name__ == '__main__':
    main()
